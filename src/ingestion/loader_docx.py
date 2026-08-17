"""Extração de texto de DOCX preservando estrutura de seções (Heading 1/2/3) e tabelas."""
from pathlib import Path

from docx import Document


def _is_heading(style_name: str) -> bool:
    return style_name.lower().startswith("heading")


def _table_to_text(table, table_index: int) -> dict:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return None
    header, *data_rows = rows
    lines = [" | ".join(header)]
    for row in data_rows:
        lines.append(" | ".join(f"{h}: {v}" for h, v in zip(header, row)))
    return {
        "text": "\n".join(lines),
        "location": f"tabela {table_index}",
        "section": "N/A",
    }


def load(file_path: Path) -> list[dict]:
    doc = Document(str(file_path))
    units: list[dict] = []

    current_section = "N/A"
    buffer: list[str] = []

    def flush():
        text = "\n".join(buffer).strip()
        if text:
            units.append({"text": text, "location": f"seção: {current_section}", "section": current_section})
        buffer.clear()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if _is_heading(para.style.name):
            flush()
            current_section = text
        else:
            buffer.append(text)
    flush()

    for i, table in enumerate(doc.tables, start=1):
        unit = _table_to_text(table, i)
        if unit:
            units.append(unit)

    return units
